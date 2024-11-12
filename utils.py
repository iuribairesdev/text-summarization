from reportlab.pdfgen import canvas
from flask import send_file
from io import BytesIO
from docx import Document
import os, glob, json

# Allowed Extensions
ALLOWED_EXTENSIONS = {'pdf'}

SETTINGS_FILE = 'settings.json'
# Define the folder to save uploaded files
UPLOAD_FOLDER = './uploaded_files'


def delete_file(file_path):
    try:
        os.remove(file_path)
        print(f"{file_path} has been deleted successfully.")
    except FileNotFoundError:
        print(f"{file_path} does not exist.")
    except PermissionError:
        print(f"Permission denied: Cannot delete {file_path}.")
    except Exception as e:
        print(f"An error occurred: {e}")


def delete_files():
    try:
        # Delete PDF files (never store PDF files)
        files = glob.glob('./uploaded_files/*')
        for f in files:
            os.remove(f)
    except FileNotFoundError:
        print(f"{f} does not exist.")
    except PermissionError:
        print(f"Permission denied: Cannot delete {f}.")
    except Exception as e:
        print(f"An error occurred: {e}")

    with open(SETTINGS_FILE, 'r') as file:
        settings = json.load(file)
    
    # if param store_p is no True
    if not settings['store_p']:    
        try:
            # Delete pickle objects
            files = glob.glob('./objects/*')
            for f in files:
                os.remove(f)
        except FileNotFoundError:
            print(f"{f} does not exist.")
        except PermissionError:
            print(f"Permission denied: Cannot delete {f}.")
        except Exception as e:
            print(f"An error occurred: {e}")


# Create a helper function to check file extension
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to validate if the input is a valid list of strings
def validate_string(str1):
    if not isinstance(str1, str):
        return False, []
    # Convert comma-separated string to a list and strip any extra spaces
    str_list = [s.strip() for s in str1.split(',') if s.strip()]
    # If the list is empty, return False
    if len(str_list) == 0:
        return False, []
    return True, str_list


# Function to export text to a PDF file
def export_text_to_pdf(text, filename):
    # Create a BytesIO buffer to hold the PDF data
    pdf_buffer = BytesIO()
    
    # Create a canvas using reportlab
    pdf_canvas = canvas.Canvas(pdf_buffer)
    pdf_canvas.setTitle(filename)

    x_position = 120
    y_position = 800

    # Split the text into lines
    lines = text.split('\n')
    for line in lines:
        pdf_canvas.drawString(x_position, y_position, line)
        y_position -= 15

        # Create a new page if the text goes too low
        if y_position < 50:
            pdf_canvas.showPage()
            y_position = 800

    pdf_canvas.save()
    print(f"Text successfully exported to {filename}")

    # Move the buffer cursor to the beginning
    pdf_buffer.seek(0)

    # Return the PDF file as an HTTP response
    return send_file(
        pdf_buffer,
        as_attachment=True,  # Set to False if you want to view in the browser
        download_name="summary_result.pdf",
        mimetype="application/pdf"
    )







def export_text_to_docx(text, filename):
   # Create a new RTF Document
    doc = Document()

    # Add a title (optional)
    doc.add_heading("Result Summary", level=1)

    # Add text content
    doc.add_paragraph(text)

    # Save the document to memory (using BytesIO)
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Go to the beginning of the byte stream


    # Return the PDF file as an HTTP response
    return send_file(
        byte_stream,
        as_attachment=True,  # Set to False if you want to view in the browser
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
