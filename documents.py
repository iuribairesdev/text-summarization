# document.py
import os
import re
import json 
from flask import flash, session, request, redirect, render_template, url_for

from flask import send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement


from auth import is_logged_in

import pandas as pd

DOCUMENTS_FILE = './documents'

def parse_text_to_dataframe(text):
    # Split the content into rows using the pipe delimiter
    rows = re.findall(r'\|(.*?)\|', text, re.DOTALL)

    # Clean and split rows into columns based on newlines
    data = [list(map(str.strip, row.split('\n'))) for row in rows]

    # Create DataFrame
    return pd.DataFrame(data[1:], columns=data[0])  # Use the first row as column headers


def export_docx(filename, pretext, df="", posttext="", table_p="0"):
   # Create a Word document
    doc = Document()

    # Add introductory paragraph
    intro_paragraph = doc.add_paragraph(pretext)
    intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if table_p == "1":
        # Add a title to the document
        title = doc.add_heading("Table Contents", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Set header row
        header_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            header_cells[i].text = column_name
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add rows to the table
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Adjust column widths (optional, if needed)
        def set_column_width(column, width):
            for cell in column.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set('w:w', str(width))
                tcW.set('w:type', 'dxa')
                tcPr.append(tcW)

        # Example: Set widths for each column (customize as needed)
        # for i, width in enumerate([2000, 1000, 5000]):  # Widths in twips (1/20 of a point)
        #     set_column_width(table.columns[i], width)



    # Add a concluding paragraph
    conclusion_paragraph = doc.add_paragraph(" ")
    conclusion_paragraph = doc.add_paragraph(" ")
    conclusion_paragraph = doc.add_paragraph(posttext)
    conclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save the document
    # output_path = './output_table.docx'
    # doc.save(output_path)
    # print(f"Document saved as {output_path}")


    # Save the document to memory (using BytesIO)
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Go to the beginning of the byte stream
    # Return the PDF file as an HTTP response
    return send_file(
        byte_stream,
        as_attachment=True,  # Set to False if you want to view in the browser
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



def export_text(text, filename, table_p):
    if table_p == "1":
        # Step 1: Find the first and last pipe positions
        start_table = text.find('|')
        end_table = text.rfind('|')
        
        if int(start_table) == 0:
            pretext = ''
            table_text = text
            posttext = ''
        else:
            # Step 2: Extract the pretext, table, and posttext
            pretext = text[:start_table-1].strip()  # Text before the first pipe
            table_text = text[start_table:end_table].strip()  # Text between the pipes
            posttext = text[end_table+1:].strip()  # Text after the last pipe

        df = parse_text_to_dataframe(table_text)

    else:
        pretext = text
        df = ''
        posttext = ''

    return export_docx(filename, pretext, df, posttext, table_p)




# Function to create a new prompt
def create_document(filename, pretext, posttext, summary):
    if request.method == 'POST':
        # print(filename, pretext, posttext)
        if filename and pretext and posttext:
            documents = read_documents()
            document_id = len(documents) + 1
            user_id = session['user']
            new_document = {'id': document_id, 'user_id': user_id, 'filename': filename, 'pretext': pretext, 'posttext': posttext, 'summary': summary}
            write_documents(new_document)
            flash('Documents created successfully!')
            return document_id
        else:
            flash('Document text cannot be empty.')
    return 


def write_documents(document):
    path = os.path.join(DOCUMENTS_FILE, session['user'])
    if not os.path.exists(path):
        os.makedirs(path)   

    # print('path', path)
    # print('document', f"{os.path.join(path, document['filename'])}.json")
    with open(f"{os.path.join(path, document['filename'])}.json", 'w') as file:
        json.dump(document, file, indent=4)

# Helper function to read prompts from JSON file
def read_documents():
    path = os.path.join(DOCUMENTS_FILE, session['user'])
    results = []
    if not os.path.exists(path):
        return []
    for filename in os.listdir(path):
        if filename.endswith('.json'):
            file_path = os.path.join(path, filename)
            with open(file_path, 'r') as file:
                try:
                    # load json data
                    data = json.load(file)
                    # extract id and filename if they exist 
                    if 'id' in data and 'filename' in data:
                        results.append({
                            'id': data['id'],
                            'filename': data['filename']
                        })
                except json.JSONDecodeError as e:
                    print(f"Error decoding JSON from file {filename}: {e}")
    return results


# Function for prompt page
def documents_page():
    # If user is not logged in, redirect to login page
    if not is_logged_in():
        return redirect(url_for('login'))
    session.pop('_flashes', None)
    # print(read_documents())
    return render_template('documents.html', documents=read_documents())




def replace_pipe_with_line_break(data):
    # This function will traverse the data and replace pipes with line breaks in strings
    if isinstance(data, dict):
        # Iterate over dictionary items
        return {key: replace_pipe_with_line_break(value) for key, value in data.items()}
    elif isinstance(data, list):
        # Iterate over list elements
        return [replace_pipe_with_line_break(element) for element in data]
    elif isinstance(data, str):
        # Replace pipes with line breaks in strings
        return data.replace(' | ', '\n')
    else:
        # If it's neither a list, dict, nor str, return as is
        return data

def document_preview(filename):
    # If user is not logged in, redirect to login page
    if not is_logged_in():
        return redirect(url_for('login'))
    
    
    if request.method == 'POST':
        filename = request.form['filename']
   
        if 'download' in request.form:
            summary = request.form['summary']
            # export to summary
            return export_text(summary, "result_summary.docx")
        elif 'cancel' in request.form:
            # Go back to the form
            return redirect(url_for('home'))
        

    file_path = f"{os.path.join(DOCUMENTS_FILE, session['user'], filename)}.json"
    with open(file_path, 'r') as file:
        # load json data
        json_data = json.load(file)
    
    # Apply pipe replacement
    document = replace_pipe_with_line_break(json_data)
               
    if not document:
        flash('Document not found.')
        return redirect(url_for('documents'))
    
    return render_template('result.html', page_title="My Documents (Archieved)", filename=filename, result=document['summary'])
